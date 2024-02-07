import os
from pdf2docx import Converter

print("Prepare to convert .pdf to .docx \n\t Just wait...")

file = 'C:/avs/tgen/python/temp/rename_file.txt'
output = "C:/avs/tgen/out"
read_file = open(file, 'r')
in_file = read_file.read()
read_file.close()
source_file = output + '/' +  in_file.removesuffix('\n') + '.pdf'

pdf_file = source_file
word_file = source_file[:-4] + '.docx'

# Convertor
cv = Converter(pdf_file)
cv.convert(word_file, start=0, end=None)
cv.close()
print('{} compiling file!'.format(source_file))


'''
import os
import sys
import aspose.words as aw
import pd
print("Prepare to convert .pdf to .docx \n Just wait...")

file = 'C:/avs/tgen/out/rename_file.txt'
output = os.path.dirname(file)
read_file = open(file, 'r')
in_file = read_file.read()
read_file.close()
source_file = output + '/' +  in_file.removesuffix('\n') + '.pdf'

doc = aw.Document(source_file)

para = doc.get_child_nodes(aw.NodeType.PARAGRAPH, True)[0].as_paragraph()
field = para.append_field(aw.fields.FieldType.FIELD_ADVANCE, False).as_field_advance()

field.down_offset = '2'
field.left_offset = '2.48'
field.right_offset = '1.1'
field.up_offset = '1.75'
field.horizontal_position = '100'
field.vertical_position = '100'
field.update()

doc.save(source_file[:-4] + '.docx')

print('{} compiling done!'.format(source_file))

'''

'''
import docx

docum = docx.Document('C:/avs/tgen/out/20gk_poln.docx')
docum.save('C:/avs/tgen/out/20gk_poln.doc')
'''

'''
import os
import aspose.pdf as pdf

# licence = pdf.License()
# licence.set_license("Aspose.Total.lic")

infile_path = os.path.abspath('C:/avs/tgen/out/KPA_205_2020-12-12_18-18-18.pdf')
infile = os.path.basename(infile_path)
document = pdf.Document(infile_path)

docSaveOptions = pdf.DocSaveOptions()

document.save(infile_path[:-4] + '.doc', docSaveOptions)

print('Rendering process completed')
'''

'''
import aspose.pdf as ap

input_pdf = 'C:/avs/tgen/out/20gk_poln.pdf'
output_pdf = 'C:/avs/tgen/out/20gk_poln.doc'

document = ap.Document(input_pdf)

save_options = ap.DocSaveOptions()
save_options.format = ap.DocSaveOptions.DocFormat.DOC_X
save_options.mode = ap.DocSaveOptions.RecognitionMode.ENHANCED_FLOW
save_options.extract_ocr_sublayer_only = True
save_options.relative_horizontal_proximity = 2.5
save_options.recognize_bullets = True

document.save(output_pdf, save_options)
'''